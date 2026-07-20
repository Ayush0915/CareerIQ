import os
import re
from utils.text_cleaner import clean_text

# Try pdfplumber first (better for multi-column resumes), fallback to PyPDF2
try:
    import pdfplumber
    _HAS_PDFPLUMBER = True
except ImportError:
    _HAS_PDFPLUMBER = False

try:
    from PyPDF2 import PdfReader
    _HAS_PYPDF2 = True
except ImportError:
    _HAS_PYPDF2 = False

try:
    import docx as _docx
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False


MAX_FILE_SIZE_MB = 5
MAX_FILE_SIZE_BYTES = MAX_FILE_SIZE_MB * 1024 * 1024


def _extract_pdf_pdfplumber(file_path: str) -> str:
    """Extract text from PDF using pdfplumber (handles columns / tables better)."""
    text_parts = []
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text(x_tolerance=3, y_tolerance=3)
                if page_text:
                    text_parts.append(page_text)
    except Exception as e:
        print(f"[parser] pdfplumber error: {e}")
    return "\n".join(text_parts)


def _extract_pdf_pypdf2(file_path: str) -> str:
    """Fallback: extract text using PyPDF2."""
    text = ""
    try:
        reader = PdfReader(file_path)
        for page in reader.pages:
            extracted = page.extract_text()
            if extracted:
                text += extracted + "\n"
    except Exception as e:
        print(f"[parser] PyPDF2 error: {e}")
    return text


def extract_text_from_pdf(file_path: str) -> str:
    """Try pdfplumber; fall back to PyPDF2."""
    text = ""
    if _HAS_PDFPLUMBER:
        text = _extract_pdf_pdfplumber(file_path)
    if not text.strip() and _HAS_PYPDF2:
        text = _extract_pdf_pypdf2(file_path)
    return text


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX, including tables."""
    if not _HAS_DOCX:
        raise RuntimeError("python-docx not installed")
    text_parts = []
    try:
        doc = _docx.Document(file_path)
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        # Also extract table cells
        for table in doc.tables:
            for row in table.rows:
                row_text = "  ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
    except Exception as e:
        print(f"[parser] DOCX error: {e}")
    return "\n".join(text_parts)


def extract_contact_info(raw_text: str) -> dict:
    """Best-effort contact field extraction — does not store anything."""
    text = raw_text[:1500]   # only scan header region

    email_match = re.search(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}", text)
    phone_match = re.search(r"(\+?\d[\d\s\-().]{7,}\d)", text)
    linkedin_match = re.search(r"linkedin\.com/in/([\w\-]+)", text, re.I)
    github_match   = re.search(r"github\.com/([\w\-]+)", text, re.I)
    name_line = text.strip().splitlines()[0] if text.strip() else ""

    return {
        "name":     name_line[:60] if len(name_line) < 60 else "",
        "email":    email_match.group(0)     if email_match    else "",
        "phone":    phone_match.group(0)     if phone_match    else "",
        "linkedin": linkedin_match.group(1)  if linkedin_match else "",
        "github":   github_match.group(1)    if github_match   else "",
    }


def validate_file_size(file_path: str) -> None:
    """Raise ValueError if file exceeds size limit."""
    size = os.path.getsize(file_path)
    if size > MAX_FILE_SIZE_BYTES:
        raise ValueError(
            f"File size {size / 1024 / 1024:.1f} MB exceeds limit of {MAX_FILE_SIZE_MB} MB."
        )


def parse_resume(file_path: str) -> dict:
    """Main resume parsing entry point."""
    validate_file_size(file_path)

    ext = file_path.rsplit(".", 1)[-1].lower()

    if ext == "pdf":
        raw_text = extract_text_from_pdf(file_path)
    elif ext == "docx":
        raw_text = extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file format: .{ext}")

    if not raw_text.strip():
        raise ValueError("Could not extract any text from the file. The PDF may be image-based or corrupted.")

    contact_info = extract_contact_info(raw_text)
    cleaned      = clean_text(raw_text)

    return {
        "raw_text":    raw_text,
        "clean_text":  cleaned,
        "contact_info": contact_info,
        "word_count":  len(cleaned.split()),
        "char_count":  len(raw_text),
    }
